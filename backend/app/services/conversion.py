"""
DocuFlow Conversion Engine
All conversion logic lives here. Each handler is a simple function.

Supported matrix:
  PDF  -> DOCX, TXT
  DOCX -> PDF, TXT, HTML
  TXT  -> DOCX, PDF
  MD   -> PDF, DOCX, HTML
  HTML -> PDF, DOCX
  Images (jpg/png/webp/heic/svg/gif) -> jpg/png/webp
  CSV  <-> XLSX
"""
import os
import shutil
import subprocess
import tempfile
from pathlib import Path

from loguru import logger


# ─── MIME type map ────────────────────────────────────────────────────────────

CONTENT_TYPES: dict[str, str] = {
    "pdf":  "application/pdf",
    "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "txt":  "text/plain",
    "html": "text/html",
    "md":   "text/markdown",
    "rtf":  "application/rtf",
    "csv":  "text/csv",
    "xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    "jpg":  "image/jpeg",
    "jpeg": "image/jpeg",
    "png":  "image/png",
    "webp": "image/webp",
    "heic": "image/heic",
    "svg":  "image/svg+xml",
    "gif":  "image/gif",
}

# Supported conversion matrix — (from_ext, to_ext) must be in this set
CONVERSION_MAP: set[tuple[str, str]] = {
    # Documents
    ("pdf",  "docx"), ("pdf",  "txt"),
    ("docx", "pdf"),  ("docx", "txt"),  ("docx", "html"),
    ("txt",  "docx"), ("txt",  "pdf"),
    ("md",   "pdf"),  ("md",   "docx"), ("md",   "html"),
    ("html", "pdf"),  ("html", "docx"),
    # Images
    ("jpg",  "png"),  ("jpg",  "webp"),
    ("jpeg", "png"),  ("jpeg", "webp"),
    ("png",  "jpg"),  ("png",  "webp"),
    ("webp", "jpg"),  ("webp", "png"),
    ("heic", "jpg"),  ("heic", "png"),
    ("svg",  "png"),
    ("gif",  "png"),  ("gif",  "jpg"),
    # Spreadsheets
    ("csv",  "xlsx"), ("xlsx", "csv"),
}


def is_supported(from_fmt: str, to_fmt: str) -> bool:
    return (from_fmt.lower(), to_fmt.lower()) in CONVERSION_MAP


def convert_file(input_path: str, from_fmt: str, to_fmt: str, output_path: str) -> str:
    """
    Convert input_path from from_fmt to to_fmt and write result to output_path.
    Returns output_path on success. Raises on failure.
    """
    f = from_fmt.lower()
    t = to_fmt.lower()

    if not is_supported(f, t):
        raise ValueError(f"Conversion .{f} → .{t} is not supported.")

    # ── PDF ──────────────────────────────────────────────
    if f == "pdf"  and t == "docx":  return _pdf_to_docx(input_path, output_path)
    if f == "pdf"  and t == "txt":   return _pdf_to_txt(input_path, output_path)

    # ── DOCX ─────────────────────────────────────────────
    if f == "docx" and t == "pdf":   return _docx_to_pdf(input_path, output_path)
    if f == "docx" and t == "txt":   return _docx_to_txt(input_path, output_path)
    if f == "docx" and t == "html":  return _docx_to_html(input_path, output_path)

    # ── TXT ──────────────────────────────────────────────
    if f == "txt"  and t == "docx":  return _txt_to_docx(input_path, output_path)
    if f == "txt"  and t == "pdf":   return _txt_to_pdf(input_path, output_path)

    # ── Markdown ─────────────────────────────────────────
    if f == "md"   and t == "html":  return _md_to_html(input_path, output_path)
    if f == "md"   and t == "pdf":   return _md_to_pdf(input_path, output_path)
    if f == "md"   and t == "docx":  return _md_to_docx(input_path, output_path)

    # ── HTML ─────────────────────────────────────────────
    if f == "html" and t == "pdf":   return _html_to_pdf(input_path, output_path)
    if f == "html" and t == "docx":  return _html_to_docx(input_path, output_path)

    # ── Images ────────────────────────────────────────────
    if f in ("jpg", "jpeg", "png", "webp", "heic", "gif") and t in ("jpg", "jpeg", "png", "webp"):
        return _convert_image(input_path, output_path, t)
    if f == "svg"  and t == "png":   return _svg_to_png(input_path, output_path)
    if f == "gif"  and t == "jpg":   return _convert_image(input_path, output_path, t)

    # ── Spreadsheets ──────────────────────────────────────
    if f == "csv"  and t == "xlsx":  return _csv_to_xlsx(input_path, output_path)
    if f == "xlsx" and t == "csv":   return _xlsx_to_csv(input_path, output_path)

    raise ValueError(f"No handler implemented for .{f} → .{t}")


# ─── PDF ──────────────────────────────────────────────────────────────────────

def _pdf_to_docx(inp: str, out: str) -> str:
    from pdf2docx import Converter
    cv = Converter(inp)
    cv.convert(out, start=0, end=None)
    cv.close()
    return out


def _pdf_to_txt(inp: str, out: str) -> str:
    import pdfplumber
    with pdfplumber.open(inp) as pdf:
        text = "\n\n".join(p.extract_text() or "" for p in pdf.pages)
    Path(out).write_text(text, encoding="utf-8")
    return out


# ─── DOCX ─────────────────────────────────────────────────────────────────────

def _docx_to_pdf(inp: str, out: str) -> str:
    """Prefer LibreOffice headless; fall back to python-docx + reportlab."""
    out_dir = str(Path(out).parent)
    try:
        result = subprocess.run(
            ["libreoffice", "--headless", "--convert-to", "pdf", "--outdir", out_dir, inp],
            capture_output=True, text=True, timeout=60,
        )
        if result.returncode == 0:
            stem = Path(inp).stem
            lo_out = Path(out_dir) / f"{stem}.pdf"
            if str(lo_out) != out:
                shutil.move(str(lo_out), out)
            return out
        logger.warning(f"LibreOffice error: {result.stderr}")
    except FileNotFoundError:
        logger.info("LibreOffice not found — using reportlab fallback")
    except Exception as e:
        logger.warning(f"LibreOffice failed: {e}")

    return _docx_to_pdf_fallback(inp, out)


def _docx_to_pdf_fallback(inp: str, out: str) -> str:
    from docx import Document
    doc = Document(inp)
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

    # Try reportlab first
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.styles import getSampleStyleSheet
        from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer
        pdf_doc = SimpleDocTemplate(out, pagesize=letter)
        styles  = getSampleStyleSheet()
        story   = []
        for text in paragraphs:
            try:
                story.append(Paragraph(text, styles["Normal"]))
                story.append(Spacer(1, 6))
            except Exception:
                pass
        if not story:
            story.append(Paragraph("(empty document)", styles["Normal"]))
        pdf_doc.build(story)
        return out
    except ImportError:
        pass

    # Fall back to fpdf2 (pure Python, always works on Windows)
    try:
        from fpdf import FPDF
        pdf = FPDF()
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.add_page()
        pdf.set_font("Helvetica", size=11)
        for text in (paragraphs or ["(empty document)"]):
            safe = text.encode("latin-1", errors="replace").decode("latin-1")
            pdf.multi_cell(0, 6, safe)
            pdf.ln(2)
        pdf.output(out)
        return out
    except ImportError:
        pass

    raise RuntimeError(
        "DOCX→PDF requires LibreOffice, reportlab, or fpdf2. "
        "Run: pip install fpdf2"
    )


def _docx_to_txt(inp: str, out: str) -> str:
    from docx import Document
    doc = Document(inp)
    text = "\n".join(p.text for p in doc.paragraphs)
    Path(out).write_text(text, encoding="utf-8")
    return out


def _docx_to_html(inp: str, out: str) -> str:
    from docx import Document
    doc = Document(inp)
    parts = [
        "<!DOCTYPE html><html><head><meta charset='utf-8'>",
        "<style>body{font-family:Georgia,serif;max-width:800px;margin:auto;padding:2rem;line-height:1.6}"
        "h1,h2,h3{margin-top:1.5em}</style></head><body>",
    ]
    for para in doc.paragraphs:
        if not para.text.strip():
            continue
        style = para.style.name or ""
        if style.startswith("Heading"):
            lvl = style[-1] if style[-1].isdigit() else "2"
            parts.append(f"<h{lvl}>{para.text}</h{lvl}>")
        else:
            parts.append(f"<p>{para.text}</p>")
    parts.append("</body></html>")
    Path(out).write_text("\n".join(parts), encoding="utf-8")
    return out


# ─── TXT ──────────────────────────────────────────────────────────────────────

def _txt_to_docx(inp: str, out: str) -> str:
    from docx import Document
    content = Path(inp).read_text(encoding="utf-8", errors="replace")
    doc = Document()
    for line in content.split("\n"):
        doc.add_paragraph(line)
    doc.save(out)
    return out


def _txt_to_pdf(inp: str, out: str) -> str:
    content = Path(inp).read_text(encoding="utf-8", errors="replace")
    html = (
        "<!DOCTYPE html><html><head><meta charset='utf-8'>"
        "<style>body{font-family:monospace;white-space:pre-wrap;padding:2rem;line-height:1.5}"
        "</style></head><body>" + content + "</body></html>"
    )
    with tempfile.NamedTemporaryFile(suffix=".html", delete=False, mode="w", encoding="utf-8") as tf:
        tf.write(html)
        tmp = tf.name
    try:
        _html_to_pdf(tmp, out)
    finally:
        os.unlink(tmp)
    return out


# ─── Markdown ─────────────────────────────────────────────────────────────────

def _md_to_html(inp: str, out: str) -> str:
    import markdown
    md_text = Path(inp).read_text(encoding="utf-8")
    body = markdown.markdown(md_text, extensions=["tables", "fenced_code", "toc"])
    html = (
        "<!DOCTYPE html><html><head><meta charset='utf-8'>"
        "<style>body{font-family:Georgia,serif;max-width:800px;margin:auto;padding:2rem;line-height:1.7}"
        "code{background:#f4f4f4;padding:2px 6px;border-radius:3px}"
        "pre{background:#f4f4f4;padding:1rem;overflow-x:auto}"
        "table{border-collapse:collapse;width:100%}th,td{border:1px solid #ddd;padding:8px}"
        "</style></head><body>" + body + "</body></html>"
    )
    Path(out).write_text(html, encoding="utf-8")
    return out


def _md_to_pdf(inp: str, out: str) -> str:
    with tempfile.NamedTemporaryFile(suffix=".html", delete=False) as tf:
        tmp = tf.name
    try:
        _md_to_html(inp, tmp)
        _html_to_pdf(tmp, out)
    finally:
        os.unlink(tmp)
    return out


def _md_to_docx(inp: str, out: str) -> str:
    with tempfile.NamedTemporaryFile(suffix=".html", delete=False) as tf:
        tmp = tf.name
    try:
        _md_to_html(inp, tmp)
        _html_to_docx(tmp, out)
    finally:
        os.unlink(tmp)
    return out


# ─── HTML ─────────────────────────────────────────────────────────────────────

def _html_to_pdf(inp: str, out: str) -> str:
    """
    Convert HTML → PDF.
    Strategy (tried in order):
      1. WeasyPrint  — best quality, but needs GTK (Linux/Mac only)
      2. wkhtmltopdf — good quality, needs binary installed
      3. fpdf2       — pure Python, Windows-safe, plain text only (fallback)
    """
    # 1. WeasyPrint (Linux/Mac with GTK)
    try:
        import weasyprint
        weasyprint.HTML(filename=inp).write_pdf(out)
        return out
    except Exception:
        pass

    # 2. wkhtmltopdf binary
    try:
        result = subprocess.run(
            ["wkhtmltopdf", "--quiet", inp, out],
            capture_output=True, timeout=60,
        )
        if result.returncode == 0 and Path(out).exists():
            return out
    except (FileNotFoundError, subprocess.TimeoutExpired):
        pass

    # 3. fpdf2 — pure Python, always works, strips HTML tags
    try:
        from fpdf import FPDF
        from bs4 import BeautifulSoup
        html_content = Path(inp).read_text(encoding="utf-8", errors="replace")
        try:
            text = BeautifulSoup(html_content, "html.parser").get_text(separator="\n")
        except Exception:
            import re
            text = re.sub(r"<[^>]+>", " ", html_content)

        pdf = FPDF()
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.add_page()
        pdf.set_font("Helvetica", size=11)
        for line in text.split("\n"):
            line = line.strip()
            if not line:
                pdf.ln(4)
                continue
            # fpdf2 encode safety
            safe = line.encode("latin-1", errors="replace").decode("latin-1")
            pdf.multi_cell(0, 6, safe)
        pdf.output(out)
        return out
    except ImportError:
        pass

    raise RuntimeError(
        "PDF conversion failed. Install fpdf2 for Windows support: pip install fpdf2"
    )


def _html_to_docx(inp: str, out: str) -> str:
    html_content = Path(inp).read_text(encoding="utf-8", errors="replace")
    try:
        from bs4 import BeautifulSoup
        text = BeautifulSoup(html_content, "html.parser").get_text(separator="\n")
    except ImportError:
        import re
        text = re.sub(r"<[^>]+>", "", html_content)

    from docx import Document
    doc = Document()
    for line in text.split("\n"):
        line = line.strip()
        if line:
            doc.add_paragraph(line)
    doc.save(out)
    return out


# ─── Images ───────────────────────────────────────────────────────────────────

def _convert_image(inp: str, out: str, to_fmt: str) -> str:
    from PIL import Image

    # HEIC support
    if inp.lower().endswith(".heic"):
        try:
            import pillow_heif
            pillow_heif.register_heif_opener()
        except ImportError:
            raise RuntimeError("HEIC support requires pillow-heif: pip install pillow-heif")

    img = Image.open(inp)

    fmt = to_fmt.lower()
    if fmt in ("jpg", "jpeg"):
        if img.mode in ("RGBA", "P", "LA"):
            bg = Image.new("RGB", img.size, (255, 255, 255))
            if img.mode == "P":
                img = img.convert("RGBA")
            if img.mode in ("RGBA", "LA"):
                bg.paste(img, mask=img.split()[-1])
                img = bg
            else:
                img = img.convert("RGB")
        elif img.mode != "RGB":
            img = img.convert("RGB")
        img.save(out, "JPEG", quality=92, optimize=True)
    elif fmt == "png":
        img.save(out, "PNG", optimize=True)
    elif fmt == "webp":
        img.save(out, "WEBP", quality=90)
    else:
        img.save(out)

    return out


def _svg_to_png(inp: str, out: str) -> str:
    """
    SVG → PNG.
    Strategy:
      1. cairosvg — best quality, needs Cairo (Linux/Mac)
      2. svglib + reportlab — pure Python, works on Windows
    """
    try:
        import cairosvg
        cairosvg.svg2png(url=inp, write_to=out)
        return out
    except Exception:
        pass

    try:
        from svglib.svglib import svg2rlg
        from reportlab.graphics import renderPM
        drawing = svg2rlg(inp)
        if drawing is None:
            raise RuntimeError("svglib could not parse the SVG file")
        renderPM.drawToFile(drawing, out, fmt="PNG")
        return out
    except ImportError:
        pass

    raise RuntimeError(
        "SVG→PNG requires either cairosvg (Linux/Mac) or svglib+reportlab (Windows). "
        "Install: pip install svglib reportlab"
    )


# ─── Spreadsheets ─────────────────────────────────────────────────────────────

def _csv_to_xlsx(inp: str, out: str) -> str:
    import pandas as pd
    df = pd.read_csv(inp)
    df.to_excel(out, index=False, engine="openpyxl")
    return out


def _xlsx_to_csv(inp: str, out: str) -> str:
    import pandas as pd
    df = pd.read_excel(inp, engine="openpyxl")
    df.to_csv(out, index=False)
    return out