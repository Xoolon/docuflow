"""
DocuFlow Watermark Engine
Applies "DocuFlow Free" watermarks to output files for free-tier users.
"""
import os
import shutil
from pathlib import Path

from loguru import logger


def watermark_pdf(input_path: str, output_path: str):
    """Stamp a diagonal semi-transparent watermark onto every PDF page."""
    try:
        from pypdf import PdfWriter, PdfReader
        import io

        try:
            from reportlab.pdfgen import canvas as rl_canvas
            from reportlab.lib.pagesizes import letter
            from reportlab.lib.colors import Color

            packet = io.BytesIO()
            c = rl_canvas.Canvas(packet, pagesize=letter)
            c.setFont("Helvetica-Bold", 52)
            c.setFillColor(Color(0.7, 0.7, 0.7, alpha=0.25))
            c.translate(300, 420)
            c.rotate(45)
            c.drawCentredString(0, 0, "DocuFlow Free")
            c.save()
            packet.seek(0)

            wm_reader = PdfReader(packet)
            wm_page = wm_reader.pages[0]

            reader = PdfReader(input_path)
            writer = PdfWriter()
            for page in reader.pages:
                page.merge_page(wm_page)
                writer.add_page(page)

            with open(output_path, "wb") as f:
                writer.write(f)

        except ImportError:
            # reportlab not available — just copy file
            shutil.copy2(input_path, output_path)
            logger.warning("reportlab not installed; PDF watermark skipped")

    except Exception as e:
        logger.error(f"PDF watermark error: {e}")
        shutil.copy2(input_path, output_path)


def watermark_docx(input_path: str, output_path: str):
    """Insert a faded header line into every section of a DOCX."""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document(input_path)
        for section in doc.sections:
            header = section.header
            if not header.paragraphs:
                header.add_paragraph()
            para = header.paragraphs[0]
            para.clear()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run("⚡ DocuFlow Free Tier — upgrade to remove watermark")
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(180, 180, 180)
            run.font.italic = True

        doc.save(output_path)

    except Exception as e:
        logger.error(f"DOCX watermark error: {e}")
        shutil.copy2(input_path, output_path)


def watermark_image(input_path: str, output_path: str):
    """Overlay a semi-transparent corner label on an image."""
    try:
        from PIL import Image, ImageDraw, ImageFont

        img = Image.open(input_path).convert("RGBA")
        width, height = img.size

        overlay = Image.new("RGBA", img.size, (0, 0, 0, 0))
        draw = ImageDraw.Draw(overlay)

        text = "DocuFlow Free"
        font_size = max(14, min(width, height) // 25)
        try:
            font = ImageFont.truetype("arial.ttf", font_size)
        except OSError:
            try:
                font = ImageFont.truetype("DejaVuSans.ttf", font_size)
            except OSError:
                font = ImageFont.load_default()

        bbox = draw.textbbox((0, 0), text, font=font)
        tw, th = bbox[2] - bbox[0], bbox[3] - bbox[1]
        x, y = width - tw - 12, height - th - 12

        draw.rectangle([x - 6, y - 4, x + tw + 6, y + th + 4], fill=(0, 0, 0, 110))
        draw.text((x, y), text, font=font, fill=(255, 255, 255, 210))

        composited = Image.alpha_composite(img, overlay)

        ext = Path(output_path).suffix.lower()
        if ext in (".jpg", ".jpeg"):
            composited.convert("RGB").save(output_path, "JPEG", quality=92)
        elif ext == ".png":
            composited.save(output_path, "PNG")
        elif ext == ".webp":
            composited.save(output_path, "WEBP", quality=90)
        else:
            composited.convert("RGB").save(output_path)

    except Exception as e:
        logger.error(f"Image watermark error: {e}")
        shutil.copy2(input_path, output_path)


def apply_watermark(input_path: str, output_path: str, file_type: str):
    """Dispatch to the correct watermark function based on file extension."""
    ext = Path(input_path).suffix.lower().lstrip(".")

    if ext == "pdf":
        watermark_pdf(input_path, output_path)
    elif ext in ("docx", "doc"):
        watermark_docx(input_path, output_path)
    elif ext in ("jpg", "jpeg", "png", "webp", "gif"):
        watermark_image(input_path, output_path)
    else:
        # txt, html, csv, etc. — no watermark applicable, just copy
        shutil.copy2(input_path, output_path)